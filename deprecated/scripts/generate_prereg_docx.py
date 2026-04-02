#!/usr/bin/env python3
"""
Generate a Word document (.docx) combining discovery results with
confirmatory preregistration hypotheses for collaborator review.

Output: results/prereg_html/prereg_discovery_results.docx

Usage:
    python scripts/generate_prereg_docx.py
"""

import json
import os
import csv
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

# ── Paths ──────────────────────────────────────────────────────────────────────
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STATS = os.path.join(ROOT, "results", "stats")
OUT_DIR = os.path.join(ROOT, "results", "prereg_html")
OUT_PATH = os.path.join(OUT_DIR, "prereg_discovery_results.docx")


# ── Helper: load JSON ─────────────────────────────────────────────────────────
def load_json(fname):
    with open(os.path.join(STATS, fname), "r") as f:
        return json.load(f)


def load_csv(fname, subfolder=None):
    path = os.path.join(STATS, subfolder, fname) if subfolder else os.path.join(STATS, fname)
    with open(path, "r") as f:
        return list(csv.DictReader(f))


# ── Load all data ─────────────────────────────────────────────────────────────
h1 = load_json("h1_lmm_results.json")
h2 = load_json("h2_coupling_results.json")
h3 = load_json("h3_optimality_results.json")
model_comp = load_csv("unified_model_comparison.csv")
affect_lmm = load_csv("affect_lmm_results.csv")
vigor_pop = load_csv("mcmc_vigor_population.csv")
indep_corr = load_csv("independent_bayesian_correlations.csv")
joint_corr = load_csv("joint_correlated_correlations.csv")
metacog = load_csv("metacognition_results.csv", subfolder="paper")

# ── Extract key numbers ──────────────────────────────────────────────────────
vigor = vigor_pop[0]
mu_delta = float(vigor["mu_delta"])
mu_delta_lo = float(vigor["mu_delta_2.5"])
mu_delta_hi = float(vigor["mu_delta_97.5"])
pct_delta_pos = float(vigor["pct_delta_pos"]) * 100

# Independent correlations
def get_indep(p1, p2):
    for row in indep_corr:
        if row["param_1"] == p1 and row["param_2"] == p2:
            return float(row["r"]), float(row["p"])
    return None, None

r_beta_delta_indep, p_beta_delta_indep = get_indep("log(\u03b2)", "\u03b4")
# Fallback to column values as stored
if r_beta_delta_indep is None:
    for row in indep_corr:
        if "beta" in row["param_1"].lower() and "delta" in row["param_2"].lower():
            r_beta_delta_indep = float(row["r"])
            p_beta_delta_indep = float(row["p"])
            break

# Joint correlations
def get_joint(p1, p2):
    for row in joint_corr:
        if row["param_1"] == p1 and row["param_2"] == p2:
            return (float(row["rho_mean"]), float(row["rho_2.5"]),
                    float(row["rho_97.5"]), float(row["P_positive"]))
    return None, None, None, None

rho_beta_delta, rho_bd_lo, rho_bd_hi, rho_bd_pp = get_joint("log_beta", "delta")

# Affect LMM
def get_affect(outcome, predictor):
    for row in affect_lmm:
        if row["outcome"] == outcome and row["predictor"] == predictor:
            return float(row["beta"]), float(row["t"]), float(row["p"])
    return None, None, None

anx_S_beta, anx_S_t, anx_S_p = get_affect("anxiety", "S_probe_z")
conf_S_beta, conf_S_t, conf_S_p = get_affect("confidence", "S_probe_z")

# Metacognition
def get_metacog(param, metric):
    for row in metacog:
        if row["parameter"] == param and row["metric"] == metric:
            return float(row["r"]), float(row["p"]), float(row["p_fdr"])
    return None, None, None

delta_anx_slope_r, delta_anx_slope_p, _ = get_metacog("delta", "S_slope_anxiety")
delta_conf_slope_r, delta_conf_slope_p, _ = get_metacog("delta", "S_slope_confidence")
delta_mean_anx_r, delta_mean_anx_p, _ = get_metacog("delta", "mean_response_anxiety")
beta_anx_slope_r, beta_anx_slope_p, _ = get_metacog("beta", "S_slope_anxiety")
beta_conf_slope_r, beta_conf_slope_p, _ = get_metacog("beta", "S_slope_confidence")
k_anx_slope_r, k_anx_slope_p, _ = get_metacog("k", "S_slope_anxiety")
k_conf_slope_r, k_conf_slope_p, _ = get_metacog("k", "S_slope_confidence")


# ── Document builder helpers ──────────────────────────────────────────────────
doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# Update heading styles
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)  # Dark navy

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_blank_line(n=1):
    for _ in range(n):
        doc.add_paragraph("")


def add_italic_para(text):
    """Add a paragraph with italic text (for hypothesis statements)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    return p


def add_bold_text(paragraph, text):
    """Add bold text to an existing paragraph."""
    run = paragraph.add_run(text)
    run.bold = True
    return run


def add_para_with_bold(text_parts):
    """Add a paragraph with alternating normal/bold text.
    text_parts is a list of (text, is_bold) tuples."""
    p = doc.add_paragraph()
    for text, bold in text_parts:
        run = p.add_run(text)
        run.bold = bold
    return p


def add_key_number(label, value, note=""):
    """Add a key number line with bold value."""
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = False
    p.add_run(str(value)).bold = True
    if note:
        p.add_run(f"  ({note})")
    return p


def format_p(p_val):
    """Format p-value for display."""
    if p_val is None:
        return "N/A"
    p_val = float(p_val)
    if p_val < 1e-100:
        return "< 10\u207b\u00b3\u2070\u2070"
    elif p_val < 1e-10:
        exp = int(f"{p_val:.0e}".split("e-")[1])
        return f"< 10\u207b{exp}"
    elif p_val < 0.001:
        return f"< .001"
    elif p_val < 0.01:
        return f"= {p_val:.3f}"
    elif p_val < 0.05:
        return f"= {p_val:.3f}"
    else:
        return f"= {p_val:.3f}"


def add_styled_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


# ── TITLE PAGE ────────────────────────────────────────────────────────────────
add_blank_line(4)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Effort Reallocation Under Threat\nin Continuous Foraging")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

add_blank_line()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Discovery Results & Confirmatory Preregistration")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

add_blank_line(2)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run("Noah Okada, Ketika Garg, Toby Wise, Dean Mobbs")
run.font.size = Pt(14)

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil.add_run("California Institute of Technology")
run.italic = True
run.font.size = Pt(12)

add_blank_line()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run("March 2026")
run.font.size = Pt(12)

add_blank_line(2)

# Sample info box
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Exploratory sample: N = 293 (all hypotheses derived from this sample)\n")
run.font.size = Pt(11)
run = info.add_run("Confirmatory sample: N = 350 (collected, not analyzed)")
run.font.size = Pt(11)
run.bold = True

doc.add_page_break()


# ── PREAMBLE ──────────────────────────────────────────────────────────────────
doc.add_heading("Overview", level=1)

doc.add_paragraph(
    "This document summarizes the key empirical findings from the exploratory sample "
    "(N = 293) and presents the confirmatory hypotheses they motivate. The purpose is to "
    "provide collaborators with a single reference that shows (1) what we found, (2) what "
    "we predict will replicate, and (3) how we will test each prediction."
)

doc.add_paragraph(
    "The confirmatory sample (N = 350) has been collected via Prolific but has not been "
    "analyzed in any form. No preprocessing, model fitting, or statistical tests have been "
    "performed on these data. All hypotheses and analysis plans below are derived exclusively "
    "from the exploratory sample."
)

doc.add_paragraph(
    "The computational framework estimates four per-subject parameters from independently "
    "fit hierarchical Bayesian models:"
)

params_list = [
    ("k", "effort discounting: the fixed cost of choosing the high-effort option (from choice model). "
          "Effort enters as a binary indicator (E = 1 for high, E = 0 for low)."),
    ("\u03b2", "threat bias: how strongly danger deters choice beyond expected value (from choice model)"),
    ("\u03b1", "baseline vigor: tonic pressing rate above task demand (from vigor model)"),
    ("\u03b4", "danger mobilization: how much excess effort increases with model-derived danger, 1 \u2212 S (from vigor model)"),
    ("\u03b3", "effort demand constraint: population-level adjustment for physical demand of the chosen option (from vigor model)"),
    ("S", "survival probability: S = (1 \u2212 T) + T/(1 + \u03bbD), integrating threat probability T and distance D"),
]
for sym, desc in params_list:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(sym).bold = True
    p.add_run(f" \u2014 {desc}")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION I — Threat shifts choice, vigor, and affect
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Section I \u2014 Threat Shifts Choice, Vigor, and Affect", level=1)

# ── Discovery text ────────────────────────────────────────────────────────────
doc.add_heading("Discovery Findings", level=2)

doc.add_paragraph(
    "Participants chose the high-effort/high-reward option on 43.1% of trials, reflecting a "
    "population-level preference for the safer, low-effort option. This proportion decreased "
    "sharply with threat (68.9% at T = 0.1, 40.1% at T = 0.5, 20.5% at T = 0.9) and with "
    "distance (61.3% at D = 1, 40.4% at D = 2, 27.7% at D = 3)."
)

# Choice behavior table
doc.add_heading("P(choose high-effort) by threat and distance", level=3)
choice_table_data = [
    ["T = 0.1", "0.808", "0.692", "0.565"],
    ["T = 0.5", "0.633", "0.381", "0.188"],
    ["T = 0.9", "0.397", "0.138", "0.078"],
]
add_styled_table(
    ["Threat", "D = 1", "D = 2", "D = 3"],
    choice_table_data
)

add_blank_line()

doc.add_paragraph(
    "The most dangerous condition (T = 0.9, D = 3) reduced high-effort choice to 7.8%, a "
    "tenfold decrease from the safest condition (T = 0.1, D = 1; 80.8%)."
)

doc.add_heading("Outcomes", level=3)
doc.add_paragraph(
    "Overall escape rate was 68.3%, with strong threat dependence: 88.5% at T = 0.1, "
    "65.2% at T = 0.5, 51.3% at T = 0.9. Mean points per trial were 0.27 (SD = 1.05), "
    "reflecting the high penalty for capture (\u22125 points) relative to cookie rewards (+1 or +5)."
)

doc.add_heading("Motor vigor", level=3)
doc.add_paragraph(
    "Capacity-normalized pressing rate averaged 0.686 (SD = 0.164), roughly 69% of "
    "individual maximum. Vigor showed a modest but consistent decrease with threat: "
    "M = 0.700 at T = 0.1, M = 0.678 at T = 0.5, M = 0.675 at T = 0.9."
)

doc.add_heading("Affect ratings", level=3)
doc.add_paragraph(
    "On the 0\u20137 scale, mean anxiety was 4.40 (SD = 1.31) and mean confidence was 3.17 "
    "(SD = 1.35). Anxiety increased from 3.72 at T = 0.1 to 5.13 at T = 0.9 (+1.41 points). "
    "Confidence decreased from 3.91 at T = 0.1 to 2.49 at T = 0.9 (\u22121.42 points)."
)

doc.add_heading("Key exploratory statistics", level=3)

# H1a
h1a = h1["H1a"]
add_key_number("H1a \u2014 Choice logistic LMM: threat",
               f"\u03b2 = {h1a['threat_z']['beta']:.3f}, z = {h1a['threat_z']['z']:.1f}",
               "p < 10\u207b\u00b3\u2070\u2070")
add_key_number("H1a \u2014 Choice logistic LMM: distance",
               f"\u03b2 = {h1a['dist_z']['beta']:.3f}, z = {h1a['dist_z']['z']:.1f}",
               "p < 10\u207b\u00b3\u2070\u2070")
add_key_number("H1a \u2014 Choice logistic LMM: threat \u00d7 distance",
               f"\u03b2 = {h1a['threat_x_dist']['beta']:.3f}, z = {h1a['threat_x_dist']['z']:.1f}",
               "p < 10\u207b\u00b3\u2070\u2070")
add_key_number("H1a \u2014 Monotonicity",
               "All 6 paired t-tests significant",
               f"all d > 0.50, all p < 10\u207b\u00b9\u2075")

add_blank_line()

# H1b
h1b = h1["H1b"]
add_key_number("H1b \u2014 Excess effort LMM: threat",
               f"\u03b2 = +{h1b['threat_z']['beta']:.4f}, z = {h1b['threat_z']['z']:.1f}",
               f"p < 10\u207b\u00b3\u2070\u2070")
add_key_number("H1b \u2014 Excess effort LMM: threat \u00d7 distance",
               f"\u03b2 = {h1b['threat_x_dist']['beta']:.4f}, z = {h1b['threat_x_dist']['z']:.1f}",
               f"p < 10\u207b\u00b3\u00b2")

add_blank_line()

# H1c
h1c_anx = h1["H1c"]["anxiety"]
h1c_conf = h1["H1c"]["confidence"]
add_key_number("H1c \u2014 Anxiety LMM: threat",
               f"\u03b2 = +{h1c_anx['threat_z']['beta']:.3f}, z = {h1c_anx['threat_z']['z']:.1f}",
               f"p < 10\u207b\u2074\u2078")
add_key_number("H1c \u2014 Confidence LMM: threat",
               f"\u03b2 = {h1c_conf['threat_z']['beta']:.3f}, z = {h1c_conf['threat_z']['z']:.1f}",
               f"p < 10\u207b\u2074\u00b2")

# ── Hypothesis ────────────────────────────────────────────────────────────────
doc.add_heading("Confirmatory Hypothesis: H1", level=2)

add_italic_para(
    "H1 \u2014 Threat will reduce high-effort choice, amplify distance-driven avoidance, "
    "increase excess motor effort, and shift anxiety upward and confidence downward."
)

doc.add_heading("H1a \u2014 High-effort choice decreases with threat and distance", level=3)
doc.add_paragraph(
    "Primary tests: Logistic mixed-effects model on trial-level binary choice: "
    "choice ~ threat_z + dist_z + threat_z:dist_z + (1 | subject). "
    "\u03b2(threat) < 0, \u03b2(dist) < 0, \u03b2(threat \u00d7 dist) < 0, all p < 0.01. "
    "Plus monotonicity of P(choose high) across all adjacent threat levels within each distance "
    "(6 paired t-tests, all p < 0.01 one-tailed)."
)

doc.add_heading("H1b \u2014 Excess effort increases with threat, diminishing at far distances", level=3)
doc.add_paragraph(
    "Primary tests: Linear mixed-effects model: excess_effort ~ threat_z * dist_z + effort_chosen_z "
    "+ (1 | subject). \u03b2(threat) > 0, p < 0.05; \u03b2(threat \u00d7 dist) < 0, p < 0.05."
)

doc.add_heading("H1c \u2014 Affect shifts with threat", level=3)
doc.add_paragraph(
    "Primary tests: LMMs on probe ratings: anxiety ~ threat_z + dist_z + (1 + threat_z | subject) "
    "and confidence ~ threat_z + dist_z + (1 + threat_z | subject). "
    "\u03b2(threat) > 0 for anxiety, \u03b2(threat) < 0 for confidence, both p < 0.001."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION II — Coherent coupling and optimality
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Section II \u2014 Coherent Coupling and Optimality", level=1)

doc.add_heading("Discovery Findings", level=2)

h2a = h2["H2a"]
h2b = h2["H2b"]
h3a = h3["H3a"]
h3b = h3["H3b"]

doc.add_paragraph(
    "Participants who shifted their choices most toward safety under high threat also showed "
    "the largest increases in motor effort \u2014 a coherent \"reallocation\" strategy."
)

add_key_number("Choice-vigor coupling",
               f"r = {h2a['r']:.3f}",
               f"p {format_p(h2a['p_one_tailed'])}, N = {h2a['N']}")
add_key_number("Split-half (choice_odd, vigor_even)",
               f"r = {h2b['split1_choice_odd_vigor_even']['r']:.3f}",
               f"p {format_p(h2b['split1_choice_odd_vigor_even']['p_one_tailed'])}")
add_key_number("Split-half (choice_even, vigor_odd)",
               f"r = {h2b['split2_choice_even_vigor_odd']['r']:.3f}",
               f"p {format_p(h2b['split2_choice_even_vigor_odd']['p_one_tailed'])}")

add_blank_line()

add_key_number("Reallocation \u2192 total earnings",
               f"r = +{h3a['r']:.3f}",
               f"p {format_p(h3a['p_one_tailed'])}")
add_key_number("Error type breakdown",
               f"{h3b['pct_cautious_of_errors']:.1%} of errors were too cautious",
               f"t = {h3b['t']:.1f}, p {format_p(h3b['p_one_tailed'])}")
add_key_number("Proportion of optimal trials",
               "69.8%",
               f"{h3b['n_optimal']} optimal out of {h3b['n_optimal'] + h3b['n_suboptimal']} total")

# ── H2 Hypothesis ────────────────────────────────────────────────────────────
doc.add_heading("Confirmatory Hypothesis: H2", level=2)

add_italic_para(
    "H2 \u2014 Choice shift and vigor shift under threat will be inversely correlated across "
    "individuals. Participants who shift choices most toward safety will show the largest "
    "increase in excess effort."
)

doc.add_heading("H2a \u2014 Choice-vigor coupling", level=3)
doc.add_paragraph(
    "Primary test: Pearson r(\u0394choice, \u0394vigor) < 0, p < 0.01 one-tailed. Where "
    "\u0394choice = P(choose high | T=0.9) \u2212 P(choose high | T=0.1) and "
    "\u0394vigor = excess_effort(T=0.9) \u2212 excess_effort(T=0.1)."
)

doc.add_heading("H2b \u2014 Split-half robustness", level=3)
doc.add_paragraph(
    "Primary test: Pearson r(\u0394choice_odd, \u0394vigor_even) < 0, p < 0.05 one-tailed. "
    "Same for the reversed split."
)

# ── H3 Hypothesis ────────────────────────────────────────────────────────────
doc.add_heading("Confirmatory Hypothesis: H3", level=2)

add_italic_para(
    "H3 \u2014 The reallocation strategy will approximate the expected-value-maximizing policy."
)

doc.add_heading("H3a \u2014 Reallocation predicts earnings", level=3)
doc.add_paragraph(
    "Primary test: Pearson r(reallocation_index, total_earnings) > 0, p < 0.01 one-tailed."
)

doc.add_heading("H3b \u2014 Dominant deviation is excessive caution", level=3)
doc.add_paragraph(
    "Primary test: Among suboptimal trials, the proportion of \"too cautious\" errors "
    "exceeds 50%. One-sample t-test against 50%: t > 2.0, p < 0.05."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION III — Computational models link choice and vigor through survival
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Section III \u2014 Computational Models Link Choice and Vigor Through Survival", level=1)

doc.add_heading("Discovery Findings", level=2)

# Model comparison
doc.add_heading("Model comparison (5 models, SVI)", level=3)
doc.add_paragraph(
    "The winning model (M5: additive effort + hyperbolic survival) achieved the highest ELBO. "
    "Additive effort outperformed multiplicative by \u0394ELBO = +158; hyperbolic survival "
    "outperformed exponential by \u0394ELBO = +174; survival models massively outperformed "
    "effort-only by \u0394ELBO = +2,038."
)

# Model comparison table
mc_headers = ["Model", "ELBO", "\u0394ELBO vs. M5"]
mc_rows = []
m5_elbo = None
for row in model_comp:
    if row["model"] == "L4a_add":
        m5_elbo = float(row["ELBO"])
        break

for row in model_comp:
    elbo = float(row["ELBO"])
    d_elbo = float(row["dELBO"])
    name_map = {
        "L0_effort": "M1: Effort-only",
        "L1_threat": "M2: Linear T\u00d7D",
        "L3_add": "M3: Exp survival + mult effort",
        "L4a_hyp": "M4: Hyp survival + mult effort",
        "L4a_add": "M5: Hyp survival + add effort",
    }
    label = name_map.get(row["model"])
    if label:
        mc_rows.append([label, f"{elbo:.1f}", f"{d_elbo:+.1f}"])

# Sort by ELBO descending
mc_rows.sort(key=lambda x: float(x[1]), reverse=True)
add_styled_table(mc_headers, mc_rows)

add_blank_line()

# Affect
doc.add_heading("Survival probability predicts affect (H4c)", level=3)
add_key_number("S \u2192 anxiety",
               f"\u03b2 = {anx_S_beta:.3f}, t = {anx_S_t:.2f}",
               f"p {format_p(anx_S_p)}")
add_key_number("S \u2192 confidence",
               f"\u03b2 = +{conf_S_beta:.3f}, t = +{conf_S_t:.2f}",
               f"p {format_p(conf_S_p)}")

add_blank_line()

# Vigor HBM
doc.add_heading("Vigor hierarchical Bayesian model (H5)", level=3)
add_key_number("\u03bc_\u03b4 (population mean danger mobilization)",
               f"+{mu_delta:.3f}",
               f"95% CI: [{mu_delta_lo:.3f}, {mu_delta_hi:.3f}]")
add_key_number("% subjects with \u03b4 > 0",
               f"{pct_delta_pos:.1f}%",
               "uninformative prior centered at zero")
add_key_number("\u03b3 (effort demand constraint)",
               "\u22120.257",
               "population-level; higher demand \u2192 less excess")

add_blank_line()

# Cross-model coupling
doc.add_heading("Cross-model parameter coupling (H6)", level=3)
add_key_number("Independent: r(log(\u03b2), \u03b4)",
               f"+{r_beta_delta_indep:.3f}" if r_beta_delta_indep else "N/A",
               f"p {format_p(p_beta_delta_indep)}" if p_beta_delta_indep else "")
add_key_number("Joint model: \u03c1(\u03b2, \u03b4)",
               f"+{rho_beta_delta:.3f}" if rho_beta_delta else "N/A",
               f"95% CI: [{rho_bd_lo:.3f}, {rho_bd_hi:.3f}]" if rho_bd_lo else "")

# ── H4 Hypothesis ────────────────────────────────────────────────────────────
doc.add_heading("Confirmatory Hypothesis: H4", level=2)

add_italic_para(
    "H4 \u2014 Choices are best explained by a model in which effort enters as an additive "
    "physical cost and survival probability follows a hyperbolic function of distance. The "
    "model-derived survival probability S predicts trial-level anxiety (negatively) and "
    "confidence (positively)."
)

doc.add_paragraph(
    "The subjective value model is: SV = R\u00b7S \u2212 k\u00b7E \u2212 \u03b2\u00b7(1\u2212S), "
    "where S = (1 \u2212 T) + T/(1 + \u03bb\u00b7D). E is binary (1 for high-effort, 0 for low-effort). "
    "This forces distance to enter exclusively through the survival function S."
)

doc.add_heading("H4a \u2014 Additive effort outperforms multiplicative effort", level=3)
doc.add_paragraph(
    "Primary test: M5 (additive) achieves higher ELBO than M4 (multiplicative). \u0394ELBO > 0."
)

doc.add_heading("H4b \u2014 Hyperbolic survival outperforms exponential survival", level=3)
doc.add_paragraph(
    "Primary test: M4 (hyperbolic) achieves higher ELBO than M3 (exponential). \u0394ELBO > 0. "
    "Additional: M5 outperforms M1 (effort-only) by \u0394ELBO > 100."
)

doc.add_heading("H4c \u2014 Survival predicts trial-level affect", level=3)
doc.add_paragraph(
    "Primary tests: LMMs with random slopes: anxiety ~ S_probe_z + (1 + S_probe_z | subject), "
    "\u03b2(S) < 0 with |t| > 3.0. confidence ~ S_probe_z + (1 + S_probe_z | subject), "
    "\u03b2(S) > 0 with |t| > 3.0."
)

# ── H5 Hypothesis ────────────────────────────────────────────────────────────
doc.add_heading("Confirmatory Hypothesis: H5", level=2)

add_italic_para(
    "H5 \u2014 Model-derived danger (1 \u2212 S) drives excess motor effort at the population "
    "level, with meaningful individual variation."
)

doc.add_paragraph(
    "The vigor model is: excess_ij = \u03b1_i + \u03b4_i\u00b7(1 \u2212 S_ij) + \u03b3\u00b7E_chosen_ij + \u03b5_ij, "
    "with hierarchical Normal priors on \u03b1_i and \u03b4_i, \u03b3 population-level, "
    "fit via NumPyro NUTS."
)

doc.add_heading("H5a \u2014 Population-mean danger mobilization is positive", level=3)
doc.add_paragraph(
    "Primary tests: (1) \u03bc_\u03b4 > 0, 95% CI excludes zero. "
    "(2) Proportion of subjects with \u03b4_i > 0 exceeds 80%. "
    "Secondary: \u03b3 < 0 (effort demand constraint)."
)

# ── H6 Hypothesis ────────────────────────────────────────────────────────────
doc.add_heading("Confirmatory Hypothesis: H6", level=2)

add_italic_para(
    "H6 \u2014 Computational parameters governing the effort-danger trade-off covary across "
    "independently estimated models."
)

doc.add_heading("H6a \u2014 \u03b2-\u03b4 coupling", level=3)
doc.add_paragraph(
    "Primary test: Pearson r(log(\u03b2_choice), \u03b4_vigor) > 0, p < .001 one-tailed."
)

doc.add_heading("H6b \u2014 Coupling confirmed by joint model", level=3)
doc.add_paragraph(
    "Secondary test: Joint hierarchical model with correlated random effects "
    "[log(k_i), log(\u03b2_i), \u03b1_i, \u03b4_i] ~ MVN(\u03bc, \u03a3), "
    "\u03a9 ~ LKJCholesky(\u03b7 = 2). "
    "\u03c1(\u03b2, \u03b4) posterior 95% CI must exclude zero."
)

doc.add_heading("H6c \u2014 \u03b2 and \u03b4 jointly predict optimality", level=3)
doc.add_paragraph(
    "Primary test: OLS regression: optimality_index ~ \u03b2_z + \u03b4_z + k_z. "
    "Both \u03b2 and \u03b4 must be significant predictors (p < 0.05)."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION IV — Vigor mobilization predicts metacognitive accuracy
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Section IV \u2014 Vigor Mobilization Predicts Metacognitive Accuracy", level=1)

doc.add_heading("Discovery Findings", level=2)

doc.add_paragraph(
    "Participants whose motor effort was more danger-responsive (higher \u03b4) also showed "
    "tighter tracking of survival probability in their subjective affect ratings. This "
    "\"metacognitive-motor bridge\" links action execution to self-report accuracy."
)

# Parameter x calibration table
doc.add_heading("Parameter \u00d7 affect-S slope correlations", level=3)

calib_headers = ["Parameter", "Metric", "r", "p", "Significant?"]
calib_rows = [
    ["\u03b4", "Anxiety slope on S", f"{delta_anx_slope_r:.3f}", format_p(delta_anx_slope_p),
     "Yes" if delta_anx_slope_p and delta_anx_slope_p < 0.05 else "No"],
    ["\u03b4", "Confidence slope on S", f"+{delta_conf_slope_r:.3f}", format_p(delta_conf_slope_p),
     "Yes" if delta_conf_slope_p and delta_conf_slope_p < 0.05 else "No"],
    ["\u03b4", "Mean anxiety", f"{delta_mean_anx_r:.3f}", format_p(delta_mean_anx_p),
     "Yes" if delta_mean_anx_p and delta_mean_anx_p < 0.05 else "No"],
    ["\u03b2", "Anxiety slope on S", f"{beta_anx_slope_r:.3f}", format_p(beta_anx_slope_p),
     "Yes" if beta_anx_slope_p and beta_anx_slope_p < 0.05 else "No"],
    ["\u03b2", "Confidence slope on S", f"+{beta_conf_slope_r:.3f}", format_p(beta_conf_slope_p),
     "Yes" if beta_conf_slope_p and beta_conf_slope_p < 0.05 else "No"],
    ["k", "Anxiety slope on S", f"+{k_anx_slope_r:.3f}", format_p(k_anx_slope_p),
     "Yes" if k_anx_slope_p and k_anx_slope_p < 0.05 else "No"],
    ["k", "Confidence slope on S", f"{k_conf_slope_r:.3f}", format_p(k_conf_slope_p),
     "Yes" if k_conf_slope_p and k_conf_slope_p < 0.05 else "No"],
]
add_styled_table(calib_headers, calib_rows)

add_blank_line()

doc.add_paragraph(
    "Key pattern: The threat-responsive parameters (\u03b4 and \u03b2) predict steeper "
    "affect-S slopes, while effort sensitivity (k) shows the opposite pattern \u2014 a "
    "dissociation between threat-responsive and cost-sensitive parameters."
)

p = doc.add_paragraph()
p.add_run("\"Adaptive, not anxious\": ").bold = True
p.add_run(
    f"High-\u03b4 individuals report lower average anxiety (r = {delta_mean_anx_r:.3f}, "
    f"p {format_p(delta_mean_anx_p)}) despite stronger anxiety-S coupling. "
    "Danger mobilization is associated with better-calibrated, not simply elevated, anxiety."
)

# ── H7 Hypothesis ────────────────────────────────────────────────────────────
doc.add_heading("Confirmatory Hypothesis: H7", level=2)

add_italic_para(
    "H7 \u2014 Participants whose motor effort is more danger-responsive (higher \u03b4) will "
    "show more accurate subjective threat appraisal."
)

doc.add_heading("H7a \u2014 \u03b4 predicts affect-S slopes", level=3)
doc.add_paragraph(
    "Primary tests:\n"
    "  1. Pearson r(\u03b4, anxiety slope on S) < 0, p < .05 one-tailed.\n"
    "  2. Pearson r(\u03b4, confidence slope on S) > 0, p < .05 one-tailed.\n"
    "Secondary: r(\u03b4, mean anxiety) < 0, p < .05 one-tailed."
)

doc.add_heading("H7b \u2014 Dissociation: threat-responsive vs. effort-sensitive parameters", level=3)
doc.add_paragraph(
    "Primary tests:\n"
    "  1. Pearson r(\u03b2, anxiety slope on S) < 0, p < .05 one-tailed.\n"
    "  2. Pearson r(k, anxiety slope on S) must be non-significant (p > .05, two-tailed).\n"
    "This dissociates threat-responsive parameters from effort cost sensitivity."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# ANALYSES SECTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Analysis Plan Summary", level=1)

doc.add_paragraph(
    "All analysis scripts used for the exploratory sample will be applied without modification "
    "to the confirmatory sample. Population parameters (\u03bb, \u03c4) are re-estimated from "
    "confirmatory data. Steps are executed in order: (1) preprocessing \u2192 (2) behavioral "
    "descriptives (H1) \u2192 (3) choice model fitting (H4a, H4b \u2192 \u03bb, k, \u03b2) "
    "\u2192 (4) affect LMMs (H4c) \u2192 (5) vigor HBM with \u03bb from step 3 (H5 \u2192 "
    "\u03b1, \u03b4) \u2192 (6) coupling and optimality (H2, H3, H6) \u2192 (7) metacognitive "
    "accuracy (H7)."
)

analyses = [
    ("H1a", "Logistic mixed-effects model (BinomialBayesMixedGLM, statsmodels): "
            "choice ~ threat_z + dist_z + threat_z:dist_z + (1 | subject)"),
    ("H1b", "Linear mixed-effects model (REML, statsmodels): "
            "excess_effort ~ threat_z * dist_z + effort_chosen_z + (1 | subject)"),
    ("H1c", "Linear mixed-effects models (REML): "
            "anxiety/confidence ~ threat_z + dist_z + (1 + threat_z | subject)"),
    ("H2", "Pearson correlations of \u0394choice and \u0394vigor (full sample + split-half)"),
    ("H3a", "Pearson correlation of reallocation index with total earnings"),
    ("H3b", "One-sample t-test: proportion of cautious errors > 50%"),
    ("H4a/b", "SVI (NumPyro, 5 models) + MCMC (NUTS, 4 chains \u00d7 1000+1000). "
              "Model comparison via ELBO and/or WAIC"),
    ("H4c", "LMM: anxiety/confidence ~ S_probe_z + (1 + S_probe_z | subject)"),
    ("H5", "Hierarchical Bayesian vigor model (NumPyro NUTS): "
            "excess = \u03b1_i + \u03b4_i\u00b7(1\u2212S) + \u03b3\u00b7E_chosen + \u03b5"),
    ("H6a", "Pearson r(log(\u03b2), \u03b4) from independently fit models"),
    ("H6b", "Joint hierarchical model with LKJ prior on correlation matrix"),
    ("H6c", "OLS regression: optimality ~ \u03b2_z + \u03b4_z + k_z"),
    ("H7a", "Pearson correlations of \u03b4 with within-subject affect-S slopes"),
    ("H7b", "Pearson correlations of \u03b2 and k with affect-S slopes (dissociation test)"),
]

for hyp, desc in analyses:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{hyp}: ").bold = True
    p.add_run(desc)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# EXCLUSION CRITERIA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Exclusion Criteria", level=1)

doc.add_paragraph(
    "Exclusion criteria are applied before any model fitting. No post-hoc exclusions "
    "based on model fit quality or statistical extremity."
)

exclusions = [
    ("Incomplete task", "Did not finish all 81 trials"),
    ("Invalid calibration", "Fewer than 10 presses across calibration trials"),
    ("Implausible keypresses",
     "Max single-trial press rate > 3 SD above sample mean (automated input), "
     "or zero presses on > 50% of regular trials (disengagement)"),
    ("Invalid predator dynamics", "> 10% of trials with physically impossible predator behavior"),
    ("Low engagement", "Overall escape rate < 35% across attack trials"),
    ("Insufficient probes",
     "< 80% probe completion (< 29/36). These subjects excluded from affect analyses "
     "(H4c, H7) only."),
]

for i, (criterion, description) in enumerate(exclusions, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(f"{criterion}: ").bold = True
    p.add_run(description)


# ══════════════════════════════════════════════════════════════════════════════
# SAMPLE SIZE JUSTIFICATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Sample Size Justification", level=1)

doc.add_paragraph(
    "N = 350 recruited via Prolific. Expected N \u2248 280\u2013330 after exclusions "
    "(exploratory retention: 83.7%)."
)

doc.add_paragraph(
    "The weakest powered test is the \u03b2-\u03b4 coupling (H6a, exploratory r = +0.279); "
    "detecting r = 0.20 at \u03b1 = 0.001 one-tailed requires N > 200 (power > 0.90). "
    "All other tests are based on large exploratory effects (choice: d > 1.5; affect LMM: "
    "t > 25; vigor: 96.6% \u03b4 > 0; choice-vigor coupling r \u2248 0.2\u20130.3)."
)

power_rows = [
    ["H1 (threat shifts behavior)", "d > 1.5", "Near certainty"],
    ["H2 (coherent coupling)", "r \u2248 0.2\u20130.3", "> 0.90 at N = 300"],
    ["H3 (optimality)", "r = 0.577", "> 0.99 at N = 200"],
    ["H4a/b (model comparison)", "\u0394ELBO = 158\u2013174", "N/A (information criterion)"],
    ["H4c (affect LMMs)", "t = \u00b125.6", "> 0.99 (even with 50% reduction)"],
    ["H5 (vigor \u03bc_\u03b4)", "P(\u03bc_\u03b4 > 0) = 1.0", "> 0.99"],
    ["H6a (\u03b2-\u03b4 coupling)", "r = +0.279", "> 0.90 at \u03b1 = .001"],
    ["H7 (metacognition)", "r \u2248 \u00b10.31", "> 0.90 at N = 200"],
]
add_styled_table(
    ["Hypothesis", "Exploratory effect", "Expected power"],
    power_rows
)

add_blank_line()

# ── Final note ────────────────────────────────────────────────────────────────
doc.add_heading("Data Availability", level=1)
doc.add_paragraph(
    "The preregistration will be timestamped on AsPredicted before the confirmatory data "
    "are opened. All analysis code will be shared on GitHub and data on OSF upon acceptance."
)


# ── Save ──────────────────────────────────────────────────────────────────────
os.makedirs(OUT_DIR, exist_ok=True)
doc.save(OUT_PATH)
print(f"Document saved to: {OUT_PATH}")
print(f"File size: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")
